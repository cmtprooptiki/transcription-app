#######################################################################################################################################################################
####################################################################  Importing Packages ##############################################################################
#######################################################################################################################################################################

import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO

#######################################################################################################################################################################
########################################################################  Set the API key  ############################################################################
#######################################################################################################################################################################

# Load the API key from secrets
openai_api_key = st.secrets["openai"]["api_key"]
# Initialize the OpenAI client
client = OpenAI(api_key = openai_api_key)

#######################################################################################################################################################################
#########################################################  Function to transcribe a segment  ##########################################################################
#######################################################################################################################################################################

def transcribe_segment(file_path):
    with open(file_path, "rb") as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
        )
    return transcription.text


#######################################################################################################################################################################
#########################################################  Function to correct Transcription  #########################################################################
#######################################################################################################################################################################

def correct_transcription(transcription,desc,speakers):
    system_prompt = f"""You are a helpful assistant that will help optimize the initial transcription.
      Your task is to correct any spelling discrepancies in the transcribed text.
        Only add necessary punctuation such as periods, commas, and capitalization, and use only the context provided.
        The text is about {desc}. The number of speakers are {speakers}.
        I'd like in the output to try to separate each speaker using a '-' and different paragraphs if the speakers are more than one.
        """

    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

#######################################################################################################################################################################
#########################################################  Function to Summarize Transcription  #######################################################################
#######################################################################################################################################################################

def summarize_transcription(transcription):
    system_prompt = f"""You are a helpful assistant that will help summarize the transcription and use only the context provided.
     Just provide the Summary and do not make further questions.
     Give the summary always on the same language as the transcription.
     """
    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

#######################################################################################################################################################################
#########################################################  Function to Create a word document  #######################################################################
#######################################################################################################################################################################
def docx_creation(summary,transcription):
    #Download the Results on Word (.docx) format
    doc=Document()
    doc.add_heading("Summary",level=2)
    doc.add_paragraph(summary)
    doc.add_paragraph("")
    doc.add_heading("Full Transcription",level=2)
    doc.add_paragraph(transcription)

    # Save the document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer